from django.shortcuts import render
from rest_framework import viewsets, permissions, status, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from .models import Document, DocumentCategory, RelatedDocument
from .serializers import (
    DocumentListSerializer, 
    DocumentDetailSerializer,
    DocumentCategorySerializer,
    RelatedDocumentSerializer,
    UserSimpleSerializer
)
from django.db.models import Count, Q
import logging
from rest_framework.views import APIView
import requests
import yaml
import os
from apps.ai.services import LangChainAssistant
import markdown
import re
import io
import docx
from bs4 import BeautifulSoup, Tag
from django.conf import settings
import html

logger = logging.getLogger(__name__)

# 读取AI配置
AI_CONFIG_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'setting.yaml')
def get_ai_config():
    with open(AI_CONFIG_PATH, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    return config['ai']

class DocumentCategoryViewSet(viewsets.ModelViewSet):
    """文档分类视图集"""
    serializer_class = DocumentCategorySerializer
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name']
    ordering_fields = ['name', 'updated_at', 'created_at']

    def get_queryset(self):
        """只返回当前用户的分类"""
        return DocumentCategory.objects.filter(user=self.request.user)

    def list(self, request, *args, **kwargs):
        """获取文档分类列表，增加日志记录和异常处理"""
        try:
            queryset = self.filter_queryset(self.get_queryset())
            logger.info(f"获取文档分类列表，用户: {request.user}，查询集数量: {queryset.count()}")
            
            page = self.paginate_queryset(queryset)
            if page is not None:
                serializer = self.get_serializer(page, many=True)
                logger.info(f"返回分页分类数据，页码: {self.paginator.page.number}，数量: {len(serializer.data)}")
                return self.get_paginated_response(serializer.data)
            
            serializer = self.get_serializer(queryset, many=True)
            logger.info(f"返回所有分类数据，数量: {len(serializer.data)}, 格式: list")
            if serializer.data:
                logger.info(f"分类数据示例: {serializer.data[0] if serializer.data else '无数据'}")
            
            return Response(serializer.data)
        except Exception as e:
            logger.error(f"获取文档分类列表出错: {str(e)}", exc_info=True)
            return Response(
                {"detail": "获取文档分类列表失败，请联系管理员。"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def perform_create(self, serializer):
        """创建分类时添加当前用户"""
        serializer.save(user=self.request.user)

class DocumentViewSet(viewsets.ModelViewSet):
    """智能文档视图集"""
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['title', 'content', 'doc_type']
    ordering_fields = ['title', 'updated_at', 'created_at']
    
    def get_queryset(self):
        user = self.request.user
        queryset = Document.objects.filter(
            Q(creator=user) | Q(shared_with=user)
        ).distinct()
        
        # 添加按分类过滤的功能
        category = self.request.query_params.get('category', None)
        if category:
            try:
                logger.info(f"按分类过滤文档，分类ID: {category}, 类型: {type(category)}")
                queryset_before = queryset.count()
                queryset = queryset.filter(category_id=category)
                queryset_after = queryset.count()
                logger.info(f"分类过滤前文档数: {queryset_before}, 过滤后: {queryset_after}")
                
                # 检查是否有匹配的分类
                from .models import DocumentCategory
                if DocumentCategory.objects.filter(id=category).exists():
                    category_name = DocumentCategory.objects.get(id=category).name
                    logger.info(f"找到分类: {category_name} (ID: {category})")
                else:
                    logger.warning(f"未找到ID为 {category} 的分类")
            except ValueError as e:
                logger.warning(f"无效的分类ID: {category}, 错误: {str(e)}")
        
        return queryset
    
    def get_serializer_class(self):
        if self.action == 'list' or self.action == 'recent':
            return DocumentListSerializer
        return DocumentDetailSerializer
    
    def perform_create(self, serializer):
        """创建文档时，自动为创建者添加权限"""
        from .models import DocumentSharePermission
        
        # 先保存文档实例
        instance = serializer.save(creator=self.request.user, last_edited_by=self.request.user)
        
        # 为创建者自己创建权限记录，确保创建者始终有权访问
        DocumentSharePermission.objects.create(
            document=instance,
            user=self.request.user,
            permission='edit'  # 创建者默认有编辑权限
        )
        
        # 新创建的文档默认未共享给他人
        instance.is_shared = False
        instance.save()
        
        logger.info(f"创建文档完成: ID={instance.id}, 标题={instance.title}")
        return instance
    
    def perform_update(self, serializer):
        instance = serializer.save(last_edited_by=self.request.user)
        logger.info(f"更新文档完成: ID={instance.id}, 标题={instance.title}")
    
    def list(self, request, *args, **kwargs):
        """获取文档列表，增加日志记录"""
        queryset = self.filter_queryset(self.get_queryset())
        logger.info(f"获取文档列表，用户: {request.user}，查询集数量: {queryset.count()}")
        
        # 获取请求的页码和每页大小
        page_size_param = request.query_params.get('page_size', self.pagination_class.page_size if hasattr(self, 'pagination_class') else 10)
        try:
            page_size = int(page_size_param)
        except (TypeError, ValueError):
            page_size = 10  # 默认值
            
        page_param = request.query_params.get('page', '1')
        try:
            page_number = int(page_param)
        except (TypeError, ValueError):
            page_number = 1  # 默认值
            
        # 计算总页数
        count = queryset.count()
        total_pages = (count + page_size - 1) // page_size if page_size > 0 else 1
        
        logger.info(f"分页参数: page={page_number}, page_size={page_size}, 总记录数={count}, 总页数={total_pages}")
        
        # 检查页码是否超出范围
        if page_number > total_pages and total_pages > 0:
            logger.warning(f"请求的页码 {page_number} 超出范围(最大页码: {total_pages})，自动返回第一页")
            # 自动调整为第一页并添加警告
            paginator = self.paginate_queryset(queryset)
            serializer = self.get_serializer(paginator, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['warning'] = f"请求的页码 {page_number} 超出范围(最大页码: {total_pages})，已返回第一页数据"
            return response
        
        # 获取分页器
        paginator = self.paginate_queryset(queryset)
        if paginator is None:
            serializer = self.get_serializer(queryset, many=True)
            logger.info(f"返回所有数据，数量: {len(serializer.data)}, 格式: list")
            return Response(serializer.data)
        
        try:
            # 正常分页处理
            serializer = self.get_serializer(paginator, many=True)
            logger.info(f"返回分页数据，页码: {self.paginator.page.number}，数量: {len(serializer.data)}")
            return self.get_paginated_response(serializer.data)
        except Exception as e:
            # 检查是否是页码超出范围的问题
            if "Page is not 'last', 'first', nor can it be converted to an int" in str(e) or "Invalid page" in str(e):
                # 获取请求的页码
                requested_page = request.query_params.get('page', '1')
                logger.warning(f"请求的页码 {requested_page} 无效或超出范围，返回第一页")
                
                # 重新获取第一页
                self.paginator.page_size = int(request.query_params.get('page_size', self.paginator.page_size))
                page_data = self.paginate_queryset(queryset)
                serializer = self.get_serializer(page_data, many=True)
                response = self.get_paginated_response(serializer.data)
                
                # 添加警告信息
                response.data['warning'] = f"请求的页码 {requested_page} 无效或超出范围，已返回第一页数据"
                return response
            
            # 其他错误直接抛出
            logger.error(f"获取文档列表分页出错: {str(e)}", exc_info=True)
            raise
    
    @action(detail=False, methods=['get'])
    def recent(self, request):
        """获取最近编辑的文档"""
        user = request.user
        recent_docs = Document.objects.filter(
            Q(creator=user) | Q(shared_with=user)
        ).distinct().order_by('-updated_at')[:10]
        
        logger.info(f"获取最近文档，用户: {user}，查询集数量: {recent_docs.count()}")
        
        serializer = self.get_serializer(recent_docs, many=True)
        logger.info(f"返回最近文档数据，数量: {len(serializer.data)}, 格式: list")
        # 这里仍然返回数组格式，前端请求处理器会将其转换为 {data: [...]} 格式
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def shared(self, request):
        """获取共享的文档"""
        user = request.user
        shared_docs = Document.objects.filter(
            Q(shared_with=user) | Q(creator=user, is_shared=True)
        ).distinct()
        
        serializer = self.get_serializer(shared_docs, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def related_documents(self, request, pk=None):
        """获取相关文档列表"""
        document = self.get_object()
        related = RelatedDocument.objects.filter(source_document=document)
        serializer = RelatedDocumentSerializer(related, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """分享文档给其他用户"""
        document = self.get_object()
        shared_users = request.data.get('shared_users', [])
        
        # 确保当前用户是文档创建者
        if document.creator != request.user:
            return Response(
                {"detail": "只有文档创建者能分享文档"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            from django.contrib.auth import get_user_model
            from .models import DocumentSharePermission
            User = get_user_model()
            
            # 清理之前的权限，重新设置
            document.permissions.all().delete()
            document.shared_with.clear()
            
            for shared_user in shared_users:
                user_id = shared_user.get('user_id')
                permission = shared_user.get('permission', 'read')  # 默认为只读权限
                
                try:
                    user = User.objects.get(id=user_id)
                    # 添加到分享用户列表
                    document.shared_with.add(user)
                    # 创建权限记录
                    DocumentSharePermission.objects.create(
                        document=document,
                        user=user,
                        permission=permission
                    )
                except User.DoesNotExist:
                    continue
            
            # 更新共享状态 - 只有当共享给非创建者用户时，才标记为共享
            other_shared_users_exist = document.shared_with.exclude(id=document.creator.id).exists()
            document.is_shared = other_shared_users_exist
            document.save(update_fields=['is_shared'])
            
            return Response(
                {"detail": "文档分享成功", "is_shared": document.is_shared},
                status=status.HTTP_200_OK
            )
        except Exception as e:
            return Response(
                {"detail": f"分享文档失败: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def generate_related(self, request, pk=None):
        """生成文档的相关文档（上下文感知）"""
        document = self.get_object()
        
        # 清除现有的相关文档
        RelatedDocument.objects.filter(source_document=document).delete()
        
        # 简单实现：基于标题和类型查找相关文档
        similar_docs = Document.objects.filter(
            ~Q(id=document.id),  # 排除当前文档
            Q(doc_type=document.doc_type) |  # 相同类型
            Q(title__icontains=document.title)  # 标题包含
        )[:5]  # 最多5个相关文档
        
        # 创建相关文档关联
        for similar_doc in similar_docs:
            # 简单的相似度计算
            title_similarity = 0.5
            if document.title.lower() in similar_doc.title.lower() or similar_doc.title.lower() in document.title.lower():
                title_similarity = 0.8
            
            type_similarity = 0.3
            if document.doc_type == similar_doc.doc_type:
                type_similarity = 0.7
            
            # 综合相似度
            relevance_score = (title_similarity + type_similarity) / 2
            
            # 创建关联
            RelatedDocument.objects.create(
                source_document=document,
                related_document=similar_doc,
                relation_type="相关文档",
                relevance_score=relevance_score
            )
        
        # 返回生成的相关文档
        related = RelatedDocument.objects.filter(source_document=document)
        serializer = RelatedDocumentSerializer(related, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def expand_content(self, request, pk=None):
        """扩展文档内容 - 使用LangChain框架扩写选中的文本"""
        document = self.get_object()
        selection = request.data.get('selection', '')
        length = request.data.get('length', 'medium')  # 扩写长度：short, medium, long
        
        if not selection:
            return Response(
                {"detail": "需要提供要扩写的文本内容"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # 根据长度设置不同的扩写指令
            length_guide = {
                'short': '适度扩展内容，增加50%左右的长度',
                'medium': '充分扩展内容，大约增加一倍的长度', 
                'long': '详细扩展内容，至少增加两倍的长度，添加更多例子和细节'
            }
            
            guide = length_guide.get(length, length_guide['medium'])
            
            # 构建提示
            prompt = f"""请扩展以下文本内容，保持原文的风格和主题，但添加更多细节、例子或解释。
原文内容:
---
{selection}
---

{guide}

请直接返回扩展后的完整内容，不要包含额外的解释或说明。保持与原文一致的语言（中文/英文）。
"""
            
            # 调用AI扩写
            assistant = LangChainAssistant()
            expanded_content = assistant.chat(prompt)
            
            # 如果AI返回超时消息，则返回错误
            if '超时' in expanded_content:
                return Response(
                    {"detail": expanded_content},
                    status=status.HTTP_504_GATEWAY_TIMEOUT
                )
            
            # 记录日志
            logger.info(f"内容扩写成功，文档ID: {document.id}, 原文长度: {len(selection)}, 扩写后长度: {len(expanded_content)}")
            
            # 返回扩写后的内容
            return Response({
                "original": selection,
                "content": expanded_content
            })
            
        except Exception as e:
            logger.error(f"内容扩写失败: {str(e)}", exc_info=True)
            return Response(
                {"detail": f"内容扩写失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['post'])
    def polish_text(self, request, pk=None):
        """润色文档文本 - 使用LangChain框架润色选中的文本"""
        document = self.get_object()
        selection = request.data.get('selection', '')
        style = request.data.get('style', 'professional')  # 风格：professional, creative, concise
        
        if not selection:
            return Response(
                {"detail": "需要提供要润色的文本内容"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # 根据风格设置不同的润色指令
            style_guide = {
                'professional': '使文本更加专业、正式，提高表达的准确性和逻辑性',
                'creative': '使文本更加生动、有创意，增加修辞和表现力', 
                'concise': '使文本更加简洁、清晰，去除冗余表达，突出关键信息'
            }
            
            guide = style_guide.get(style, style_guide['professional'])
            
            # 构建提示
            prompt = f"""请润色以下文本内容，提高其质量，但保持原意不变。
原文内容:
---
{selection}
---

{guide}

请直接返回润色后的完整内容，不要包含额外的解释或说明。保持与原文一致的语言（中文/英文）。
"""
            
            # 调用AI润色
            assistant = LangChainAssistant()
            polished_content = assistant.chat(prompt)
            
            # 如果AI返回超时消息，则返回错误
            if '超时' in polished_content:
                return Response(
                    {"detail": polished_content},
                    status=status.HTTP_504_GATEWAY_TIMEOUT
                )
            
            # 记录日志
            logger.info(f"文本润色成功，文档ID: {document.id}, 风格: {style}")
            
            # 返回润色后的内容
            return Response({
                "original": selection,
                "content": polished_content
            })
            
        except Exception as e:
            logger.error(f"文本润色失败: {str(e)}", exc_info=True)
            return Response(
                {"detail": f"文本润色失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['post'])
    def grammar_check(self, request, pk=None):
        """AI语法纠错 - 检查语法、错别字、表达流畅性，返回优化建议和优化文本"""
        document = self.get_object()
        selection = request.data.get('selection', '')
        if not selection:
            return Response(
                {"detail": "需要提供要纠错的文本内容"},
                status=status.HTTP_400_BAD_REQUEST
            )
        try:
            prompt = (
                "请仔细检查以下文本的语法、错别字和表达流畅性，指出所有问题并给出详细优化建议，然后输出优化后的文本。\n\n"
                "【原文内容】：\n" + selection +
                "\n\n【输出格式要求】：\n"
                "1. 先用中文列出所有发现的问题和优化建议（如有错别字请高亮标注，语法问题请简明说明，表达不通顺请给出更优表达），每条建议单独分行。\n"
                "2. 然后输出优化后的完整文本。\n"
                "3. 不要输出与内容无关的解释。"
            )
            assistant = LangChainAssistant()
            ai_response = assistant.chat(prompt)
            # 智能分割优化建议和优化文本
            corrected = ''
            suggestions = ai_response
            # 1. 正则优先提取"优化后的完整文本"后内容
            match = re.search(r'优化后的完整文本[：:]?\s*([\s\S]+)', ai_response)
            if match:
                corrected = match.group(1).strip()
                suggestions = ai_response[:match.start()].strip()
            else:
                # 2. 尝试"2."、"2、"等分割
                match2 = re.search(r'2[\.|、|:|：]\s*([\s\S]+)', ai_response)
                if match2:
                    corrected = match2.group(1).strip()
                    suggestions = ai_response[:match2.start()].strip()
            # 3. 如果还没有，兜底用最后一段
            if not corrected:
                parts = re.split(r'\n{2,}', ai_response)
                if len(parts) > 1:
                    corrected = parts[-1].strip()
                    suggestions = '\n\n'.join(parts[:-1]).strip()
            # 4. 如果还是没有，全部给建议，优化文本为空
            # 去除corrected前的"**优化后的文本**："等前缀
            if corrected:
                corrected = re.sub(r'^([*\s]*优化后的文本[*\s]*[:：]*)', '', corrected).strip()
            return Response({
                "suggestions": suggestions,
                "corrected": corrected
            })
        except Exception as e:
            logger.error(f"AI语法纠错失败: {str(e)}", exc_info=True)
            return Response(
                {"detail": f"AI语法纠错失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def create(self, request, *args, **kwargs):
        """自定义创建文档方法，确保返回正确的响应格式"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        instance = self.perform_create(serializer)
        
        # 获取完整的序列化表示
        headers = self.get_success_headers(serializer.data)
        # 使用刚创建的实例重新获取数据，确保包含所有必要字段
        response_serializer = self.get_serializer(instance)
        
        response_data = response_serializer.data
        logger.info(f"创建文档API响应: {response_data}")
        logger.info(f"响应数据类型: {type(response_data)}")
        
        # 确保返回包含id的对象
        if isinstance(response_data, dict) and 'id' in response_data:
            logger.info(f"文档创建成功, ID: {response_data['id']}")
        else:
            logger.error(f"响应数据格式错误: {response_data}")
        
        return Response(
            response_data, 
            status=status.HTTP_201_CREATED, 
            headers=headers
        )

    @action(detail=True, methods=['post'])
    def translate(self, request, pk=None):
        """AI多语言翻译，支持选中文本或全文一键翻译为指定目标语言"""
        document = self.get_object()
        selection = request.data.get('selection', '')
        target_lang = request.data.get('target_lang', 'en')
        if not selection:
            return Response({"detail": "需要提供要翻译的文本内容"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            prompt = f"请将以下文本100%全部翻译为{target_lang}，严格只输出一次完整翻译结果，不要重复，不要输出任何解释、注释、原文或其它内容，只输出翻译后的目标语言文本：\n\n{selection}"
            assistant = LangChainAssistant()
            translated = assistant.chat(prompt)
            # 后处理：去除AI返回的注释、括号说明、解释等，仅保留翻译正文
            # 1. 去除常见的括号注释
            translated = re.sub(r'\(.*?\)', '', translated)
            translated = re.sub(r'（.*?）', '', translated)
            # 2. 去除"注：""Note:"等开头的段落
            translated = re.sub(r'(^|\n)[（(\[]?注[:：].*?(\n|$)', '\n', translated, flags=re.IGNORECASE)
            translated = re.sub(r'(^|\n)[（(\[]?note[:：].*?(\n|$)', '\n', translated, flags=re.IGNORECASE)
            # 3. 去除多余空行
            translated = re.sub(r'\n{2,}', '\n', translated)
            translated = translated.strip()
            return Response({"translated": translated})
        except Exception as e:
            logger.error(f"AI翻译失败: {e}")
            return Response({"detail": "AI翻译失败"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def qa(self, request, pk=None):
        """文档智能问答，支持基于当前文档内容和知识库即时解答"""
        document = self.get_object()
        question = request.data.get('question', '')
        context = request.data.get('context', '') or document.content
        if not question:
            return Response({"detail": "请输入您的问题"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            # 过滤上下文中可能的坐标数据
            if context:
                import re
                # 安全检查：过滤掉可能是坐标/向量的数值行
                filtered_lines = []
                for line in context.split('\n'):
                    # 跳过只包含数字、点和空格的行 (可能是坐标)
                    if re.match(r'^[\d\.\s]+$', line.strip()):
                        continue
                    # 跳过向量格式的行
                    if re.match(r'^\[[\d\.\s,]+\]$', line.strip()):
                        continue
                    filtered_lines.append(line)
                
                # 重新组合过滤后的内容
                context = '\n'.join(filtered_lines)
                
            prompt = f"请基于以下文档内容和知识库，精准、简明地回答用户问题。\n\n【文档内容】：\n{context}\n\n【用户问题】：{question}\n\n请直接输出答案，不要输出与答案无关的解释。"
            assistant = LangChainAssistant()
            answer = assistant.chat(prompt)
            return Response({"answer": answer.strip()})
        except Exception as e:
            logger.error(f"AI问答失败: {str(e)}", exc_info=True)
            return Response({"detail": f"AI问答失败: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['post'], url_path='import/word')
    def import_word(self, request):
        """导入Word文档，转换为HTML格式"""
        try:
            # 获取上传的文件
            file = request.FILES.get('file')
            if not file:
                return Response(
                    {"detail": "未提供文件"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 检查文件类型
            if not file.name.endswith(('.docx', '.doc')):
                return Response(
                    {"detail": "仅支持.docx或.doc格式"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 读取文件内容
            file_content = file.read()
            
            # 使用python-docx解析Word文档
            doc = docx.Document(io.BytesIO(file_content))
            
            # 提取文件名作为文档标题
            title = os.path.splitext(file.name)[0]
            
            # 将Word文档转换为HTML
            html_content = self._convert_word_to_html(doc)
            
            # 创建新文档
            document = Document.objects.create(
                title=title,
                content=html_content,
                doc_type='Word导入',
                creator=request.user,
                last_edited_by=request.user
            )
            
            # 返回文档详情
            serializer = DocumentDetailSerializer(document)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        
        except Exception as e:
            logger.error(f"导入Word文档失败: {str(e)}", exc_info=True)
            return Response(
                {"detail": f"导入失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _convert_word_to_html(self, doc):
        """将Word文档转换为HTML格式，适配Quill编辑器"""
        html_parts = []
        
        # 处理每个段落
        for para in doc.paragraphs:
            if not para.text.strip():
                # 空段落转换为<p><br></p>
                html_parts.append('<p><br></p>')
                continue
            
            # 段落样式处理
            if para.style.name and para.style.name.startswith('Heading'):
                # 根据标题级别设置标签
                heading_level = para.style.name[-1]
                if heading_level.isdigit() and 1 <= int(heading_level) <= 6:
                    html_parts.append(f'<h{heading_level}>{html.escape(para.text)}</h{heading_level}>')
                else:
                    # 默认标题样式
                    html_parts.append(f'<h2>{html.escape(para.text)}</h2>')
            else:
                # 处理粗体、斜体等文本样式
                styled_text = self._process_text_with_styles(para)
                html_parts.append(f'<p>{styled_text}</p>')
        
        # 处理表格
        for table in doc.tables:
            table_html = '<table style="border-collapse: collapse; width: 100%;">'
            for row in table.rows:
                table_html += '<tr>'
                for cell in row.cells:
                    # 提取单元格的文本
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if cell_text:
                            cell_text += "<br>"
                        cell_text += self._process_text_with_styles(paragraph)
                    
                    table_html += f'<td style="border: 1px solid #ddd; padding: 8px;">{cell_text}</td>'
                table_html += '</tr>'
            table_html += '</table>'
            html_parts.append(table_html)
        
        # 合并所有HTML内容
        html_content = ''.join(html_parts)
        
        return html_content
    
    def _process_text_with_styles(self, paragraph):
        """处理段落中的文本样式"""
        result = ""
        for run in paragraph.runs:
            text = html.escape(run.text)
            
            # 应用样式
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
            
            result += text
        
        return result

    @action(detail=False, methods=['get'])
    def users(self, request):
        """获取可分享文档的用户列表"""
        from django.contrib.auth import get_user_model
        User = get_user_model()
        
        # 排除当前用户自己
        users = User.objects.exclude(id=request.user.id).order_by('username')
        serializer = UserSimpleSerializer(users, many=True)
        
        return Response(serializer.data)

class AIGenerateDocumentView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        """AI生成文档接口（LangChain实现+类型识别+markdown转html+超时处理）"""
        data = request.data
        title = data.get('title')
        requirement = data.get('requirement')
        category_id = data.get('category_id')
        doc_type = data.get('doc_type', 'AI')
        if not title or not requirement:
            return Response({'detail': '文档主题和要求不能为空'}, status=400)
        prompt = f"请根据以下主题和要求生成一份详细的文档（使用markdown格式）：\n主题：{title}\n要求：{requirement}"
        try:
            assistant = LangChainAssistant()
            markdown_content = assistant.chat(prompt)
            if '超时' in markdown_content:
                return Response({'detail': markdown_content}, status=504)
            # markdown转html
            html_content = markdown.markdown(markdown_content, extensions=['extra', 'tables', 'sane_lists'])
            # 类型识别
            type_prompt = (
                "请根据以下文档内容，判断其最准确的文档类型，只返回类型名称（如：报告、总结、计划、通知、方案、合同、请示、纪要、简历、申请、说明、公告、邮件等），不要解释：\n" + markdown_content[:2000]
            )
            ai_type = assistant.chat(type_prompt)
            if '超时' in ai_type:
                ai_type = 'AI'
            # 保存文档
            doc = Document.objects.create(
                title=title,
                content=html_content,
                doc_type=ai_type.strip(),
                category_id=category_id if category_id else None,
                creator=request.user,
                last_edited_by=request.user
            )
            return Response({'id': doc.id, 'title': doc.title, 'doc_type': doc.doc_type}, status=201)
        except Exception as e:
            logger.error(f"AI生成文档失败: {e}")
            return Response({'detail': f'AI生成文档失败: {e}'}, status=500)

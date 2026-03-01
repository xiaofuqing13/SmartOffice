from rest_framework import viewsets, permissions, status, filters
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.views import APIView
from rest_framework.pagination import PageNumberPagination
from django_filters.rest_framework import DjangoFilterBackend
from django.db.models import Q
from django.utils import timezone
from django.http import HttpResponse
from io import BytesIO
import docx
import re
import html
import logging
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup

from .models import Contract, ContractVersion, ContractAttachment, ContractAction, ContractTemplate
from .serializers import (
    ContractListSerializer, ContractDetailSerializer, 
    ContractCreateSerializer, ContractUpdateSerializer,
    ContractAttachmentSerializer, ContractTemplateListSerializer,
    ContractTemplateDetailSerializer, ContractTemplateCreateUpdateSerializer
)
from .langchain_agents import polish_contract, check_contract, invalidate_polish_cache, invalidate_check_cache, query_regulation, process_html_entities
from .langchain_contract_generator import generate_contract

# 新增：用于获取选项列表的视图
class ContractTypeListView(APIView):
    """获取所有唯一的合同类型"""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        types = ContractTemplate.objects.values_list('contract_type', flat=True).distinct()
        # 过滤掉空值或None值
        filtered_types = [t for t in types if t]
        return Response([{'name': t} for t in filtered_types])

class IndustryListView(APIView):
    """获取所有唯一的行业领域"""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        industries = ContractTemplate.objects.values_list('industry', flat=True).distinct()
        filtered_industries = [i for i in industries if i]
        return Response([{'name': i} for i in filtered_industries])

class SceneListView(APIView):
    """获取所有唯一的交易场景"""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        scenes = ContractTemplate.objects.values_list('scene', flat=True).distinct()
        filtered_scenes = [s for s in scenes if s]
        return Response([{'name': s} for s in filtered_scenes])

# 获取日志记录器
logger = logging.getLogger(__name__)

# 定义合同分页类
class ContractPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100

class ContractTemplateViewSet(viewsets.ModelViewSet):
    """合同模板视图集"""
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['contract_type', 'industry', 'scene', 'is_active']
    search_fields = ['name', 'description']
    ordering_fields = ['created_at', 'name', 'contract_type']
    ordering = ['-created_at']
    
    def get_queryset(self):
        """获取模板列表，可根据查询参数筛选"""
        queryset = ContractTemplate.objects.all()
        
        # 根据合同类型筛选
        contract_type = self.request.query_params.get('contract_type', None)
        if contract_type:
            queryset = queryset.filter(contract_type=contract_type)
        
        # 根据行业筛选
        industry = self.request.query_params.get('industry', None)
        if industry:
            queryset = queryset.filter(industry=industry)
        
        # 根据场景筛选
        scene = self.request.query_params.get('scene', None)
        if scene:
            queryset = queryset.filter(scene=scene)
            
        return queryset
    
    def get_serializer_class(self):
        """根据不同操作返回不同的序列化器"""
        if self.action == 'list':
            return ContractTemplateListSerializer
        elif self.action == 'retrieve':
            return ContractTemplateDetailSerializer
        elif self.action in ['create', 'update', 'partial_update']:
            return ContractTemplateCreateUpdateSerializer
        return ContractTemplateListSerializer
    
    def perform_create(self, serializer):
        """创建模板时自动设置创建者"""
        serializer.save(created_by=self.request.user)
    
    @action(detail=True, methods=['post'])
    def use_template(self, request, pk=None):
        """使用模板创建合同"""
        template = self.get_object()
        
        # 创建合同数据 - 只使用合同必要的字段
        contract_data = {
            'title': request.data.get('title', template.name),
            'company': request.data.get('company', ''),
            'type': request.data.get('type', template.contract_type),
            'content': template.content,
            'remark': request.data.get('remark', ''),
            'template_id': template.id,  # 使用template_id而不是template对象
            'amount': request.data.get('amount', 0.00),
        }
        
        # 创建合同
        serializer = ContractCreateSerializer(data=contract_data)
        if serializer.is_valid():
            contract = serializer.save(created_by=request.user)
            
            # 创建操作记录
            ContractAction.objects.create(
                contract=contract,
                user=request.user,
                action_type='create_from_template',
                description=f'从模板"{template.name}"创建了合同'
            )
            
            # 创建成功后，清除此ID的潜在缓存
            invalidate_polish_cache(contract.id)
            invalidate_check_cache(contract.id)
            
            return Response(
                ContractDetailSerializer(contract).data,
                status=status.HTTP_201_CREATED
            )
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ContractViewSet(viewsets.ModelViewSet):
    """合同视图集"""
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_fields = ['type']
    ordering_fields = ['created_at', 'expire_date', 'amount', 'title']
    ordering = ['-created_at']
    pagination_class = ContractPagination

    def get_permissions(self):
        """根据不同的请求方法设置权限"""
        if self.action == 'download':
            # 下载操作使用更宽松的权限
            return [permissions.IsAuthenticated()]
        return super().get_permissions()

    def get_queryset(self):
        """获取合同列表，可根据查询参数筛选"""
        queryset = Contract.objects.filter(created_by=self.request.user)
        
        # 手动处理搜索, 兼容 'search' 和 'q' 两种可能的参数名
        search_query = self.request.query_params.get('search', None) or self.request.query_params.get('q', None)
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query) |
                Q(number__icontains=search_query) |
                Q(company__icontains=search_query)
            )
            
        return queryset

    def get_serializer_class(self):
        """根据不同操作返回不同的序列化器"""
        if self.action == 'list':
            return ContractListSerializer
        elif self.action == 'retrieve':
            return ContractDetailSerializer
        elif self.action == 'create':
            return ContractCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return ContractUpdateSerializer
        return ContractListSerializer

    def perform_create(self, serializer):
        """创建合同时自动设置创建者"""
        try:
            # 调试输出
            print(f"尝试创建合同. 用户: {self.request.user}, 数据: {serializer.validated_data}")
            
            # 保存并自动设置创建者
            serializer.save(created_by=self.request.user, number='')
            
            # 创建成功后，清除此ID的潜在缓存
            contract = serializer.instance
            invalidate_polish_cache(contract.id)
            invalidate_check_cache(contract.id)
            
            # 创建操作记录
            ContractAction.objects.create(
                contract=contract,
                user=self.request.user,
                action_type='create',
                description='创建了合同'
            )
        except Exception as e:
            print(f"创建合同时发生错误: {e}")
            raise

    def perform_update(self, serializer):
        """更新合同时记录操作"""
        # 保存更新前清除AI润色和检查缓存
        if 'content' in self.request.data:
            contract_id = serializer.instance.id
            invalidate_polish_cache(contract_id)
            invalidate_check_cache(contract_id)
            
        serializer.save()
        
        # 创建操作记录
        contract = serializer.instance
        ContractAction.objects.create(
            contract=contract,
            user=self.request.user,
            action_type='update',
            description=f'更新了合同信息'
        )
        
        # 如果内容发生变化，创建新版本
        if 'content' in self.request.data:
            # 获取最新版本号
            latest_version = contract.versions.order_by('-version').first()
            new_version = 1
            if latest_version:
                new_version = latest_version.version + 1
                
            # 创建新版本
            ContractVersion.objects.create(
                contract=contract,
                version=new_version,
                content=contract.content,
                created_by=self.request.user
            )

    @action(detail=True, methods=['post'])
    def add_attachment(self, request, pk=None):
        """添加合同附件"""
        contract = self.get_object()
        
        # 创建附件
        serializer = ContractAttachmentSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(
                contract=contract,
                uploaded_by=request.user
            )
            
            # 创建操作记录
            ContractAction.objects.create(
                contract=contract,
                user=request.user,
                action_type='add_attachment',
                description=f'添加了附件：{serializer.validated_data["name"]}'
            )
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def check_contract(self, request, pk=None):
        """AI检查合同风险"""
        contract = self.get_object()
        try:
            # 准备数据对象
            data = None
            
            # 从请求中获取内容和偏好设置
            if 'content' in request.data:
                # 预处理请求内容，替换HTML实体
                content = request.data['content']
                processed_content = process_html_entities(content)
                data = {'content': processed_content}
                
                # 如果有偏好设置，添加到数据对象
                if 'preferences' in request.data:
                    data['preferences'] = request.data['preferences']
                    print(f"收到合同检查偏好设置: {request.data['preferences']}")
            
            # 记录完整请求数据，便于调试
            print(f"完整合同检查请求数据: {data}")
            
            # 使用LangChain代理分析合同
            result = check_contract(pk, data)
            
            # 如果有错误，返回错误信息
            if 'error' in result:
                return Response(
                    {"error": result['error']},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 创建操作记录
            ContractAction.objects.create(
                contract=self.get_object(),
                user=request.user,
                action_type='check_contract',
                description='使用AI检查了合同内容'
            )
            
            return Response(result, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"error": f"合同文本检查失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['post'])
    def ai_polish(self, request, pk=None):
        """AI润色分析合同"""
        try:
            # 检查请求是否为实时分析 
            is_realtime = request.query_params.get('real_time', 'false').lower() == 'true'
            
            # 准备传递给后端处理函数的数据
            data = {}
            
            # 从请求中获取所有相关参数
            if 'content' in request.data:
                data['content'] = request.data['content']
            
            # 如果有is_manual标记，传递它
            if 'is_manual' in request.data:
                data['is_manual'] = request.data['is_manual']
                
            # 如果是实时请求，添加标记以便后端快速处理
            if is_realtime:
                data['is_realtime'] = True
                
            # 如果有个性化偏好设置，传递它们
            if 'preferences' in request.data:
                data['preferences'] = request.data['preferences']
                print(f"收到个性化设置: {request.data['preferences']}")
            
            # 记录完整请求数据，便于调试
            print(f"完整AI润色请求数据: {data}")
            
            # 使用LangChain代理分析合同
            result = polish_contract(pk, data)
            
            # 如果有错误，返回错误信息
            if 'error' in result:
                return Response(
                    {"error": result['error']},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 只有非实时请求才记录操作历史
            if not is_realtime:
                # 创建操作记录
                ContractAction.objects.create(
                    contract=self.get_object(),
                    user=request.user,
                    action_type='ai_polish',
                    description='使用AI智能润色分析了合同'
                )
            
            return Response(result, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"error": f"AI润色分析失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def create(self, request, *args, **kwargs):
        """创建合同，支持AI生成"""
        # 检查是否为AI生成合同请求
        use_ai_agent = request.data.get('use_ai_agent', False)
        
        if use_ai_agent:
            try:
                # 获取用户描述
                description = request.data.get('description', '')
                if not description:
                    return Response(
                        {"error": "合同描述不能为空"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
                
                # 使用AI生成合同
                contract_data = generate_contract(description)
                metadata = contract_data.get("metadata", {})
                content = contract_data.get("content", "")
                
                if not content:
                    return Response(
                        {"error": "生成合同内容失败"}, 
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
                
                # 创建合同对象
                contract = Contract.objects.create(
                    title=metadata.get("title", "AI生成的合同"),
                    type=metadata.get("type", "其他"),
                    company=metadata.get("party_b", "未指定公司"),
                    content=content,
                    amount=float(metadata.get("amount", 0)),
                    created_by=request.user
                )
                
                # 更新合同编号（在保存时自动生成）
                contract.save()
                
                # 创建成功后，清除此ID的潜在缓存
                invalidate_polish_cache(contract.id)
                invalidate_check_cache(contract.id)
                
                # 创建操作记录
                ContractAction.objects.create(
                    contract=contract,
                    user=request.user,
                    action_type='create',
                    description='使用AI生成了合同'
                )
                
                # 返回合同数据
                return Response(
                    ContractDetailSerializer(contract).data,
                    status=status.HTTP_201_CREATED
                )
            
            except Exception as e:
                import traceback
                logger.error(f"AI生成合同失败: {str(e)}")
                logger.error(traceback.format_exc())
                return Response(
                    {"error": f"生成合同失败: {str(e)}"}, 
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        
        # 正常创建合同流程
        return super().create(request, *args, **kwargs)

    @action(detail=True, methods=['get'], url_path='download-docx', url_name='download_docx')
    def download(self, request, pk=None):
        """下载合同为Word文档"""
        try:
            print(f"---------- 开始下载合同文档 ----------")
            print(f"合同ID: {pk}")
            print(f"请求参数: {request.query_params}")
            print(f"请求用户: {request.user}")
            
            # 权限检查 - 确保用户已认证
            if not request.user.is_authenticated:
                print("用户未认证，拒绝访问")
                return Response(
                    {"error": "需要身份验证才能下载合同"},
                    status=status.HTTP_401_UNAUTHORIZED
                )
                
            # 尝试获取合同对象
            try:
                contract = self.get_object()
                print(f"成功获取合同: {contract.id} - {contract.title}")
            except Exception as e:
                print(f"获取合同对象失败: {str(e)}")
                return Response(
                    {"error": f"合同不存在或无访问权限: {str(e)}"},
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # 获取请求格式，默认为docx
            file_format = request.query_params.get('format', 'docx').lower()
            print(f"下载格式: {file_format}")
            
            if file_format == 'docx':
                try:
                    # 创建docx文档
                    document = docx.Document()
                    
                    # 设置页面边距
                    sections = document.sections
                    for section in sections:
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                    
                    # 添加标题
                    title = document.add_heading(contract.title or '合同', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 添加基本信息
                    info_table = document.add_table(rows=1, cols=2)
                    info_table.style = 'Table Grid'
                    info_cells = info_table.rows[0].cells
                    info_cells[0].text = '合同编号'
                    info_cells[1].text = contract.number or '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '合同类型'
                    row.cells[1].text = contract.type or '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '签约对方'
                    row.cells[1].text = contract.company or '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '合同金额'
                    row.cells[1].text = f"¥{contract.amount}" if contract.amount else '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '签约日期'
                    row.cells[1].text = str(contract.sign_date) if contract.sign_date else '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '生效日期'
                    row.cells[1].text = str(contract.start_date) if contract.start_date else '未设置'
                    
                    row = info_table.add_row()
                    row.cells[0].text = '到期日期'
                    row.cells[1].text = str(contract.expire_date) if contract.expire_date else '未设置'
                    
                    # 添加空行
                    document.add_paragraph()
                    
                    # 添加合同内容 - 将HTML转换为纯文本
                    document.add_heading('合同内容', level=1)
                    
                    # 创建处理段落样式的函数
                    def apply_paragraph_style(paragraph, p_element):
                        # 处理对齐方式
                        if p_element.get('style'):
                            style = p_element.get('style').lower()
                            if 'text-align:center' in style or 'text-align: center' in style:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif 'text-align:right' in style or 'text-align: right' in style:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif 'text-align:justify' in style or 'text-align: justify' in style:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        # 检查class属性中是否包含对齐信息
                        if p_element.get('class'):
                            classes = ' '.join(p_element.get('class')).lower()
                            if 'text-center' in classes or 'ql-align-center' in classes:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif 'text-right' in classes or 'ql-align-right' in classes:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif 'text-justify' in classes or 'ql-align-justify' in classes:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                
                        # 如果上述方法未能检测到对齐方式，检查HTML元素本身是否有Quill特有的对齐属性
                        if paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:  # 默认值
                            if 'class="ql-align-center"' in str(p_element):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif 'class="ql-align-right"' in str(p_element):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif 'class="ql-align-justify"' in str(p_element):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # 创建处理复选框的函数
                    def handle_checkbox(paragraph, text, p_element=None):
                        # 保留原始文本中的空格模式
                        # 检测连续的空格
                        space_groups = re.findall(r' {2,}', text)
                        
                        # 先处理连续空格，将它们转换为特殊标记，避免被后续处理拆分
                        for i, spaces in enumerate(space_groups):
                            marker = f"__SPACE_GROUP_{i}__"
                            text = text.replace(spaces, marker)
                        
                        # 拆分文本，保留复选框字符
                        parts = re.split(r'(□|\u2610|\u25A1)', text)
                        for part in parts:
                            if part in ['□', '\u2610', '\u25A1']:
                                # 复选框字符 - 使用特殊字体或符号
                                checkbox_run = paragraph.add_run('□')
                                checkbox_run.font.name = 'Symbol'
                            elif part:
                                # 恢复空格组
                                for i, spaces in enumerate(space_groups):
                                    marker = f"__SPACE_GROUP_{i}__"
                                    part = part.replace(marker, spaces)
                                    
                                # 添加文本
                                paragraph.add_run(part)
                        
                        # 如果提供了原始HTML元素，应用其样式和对齐方式
                        if p_element:
                            apply_paragraph_style(paragraph, p_element)
                    
                    if contract.content:
                        # 预处理HTML内容
                        content_html = contract.content
                        
                        # 手动替换所有&nbsp;为实际空格（多次替换确保全部替换）
                        while '&nbsp;' in content_html:
                            content_html = content_html.replace('&nbsp;', ' ')
                            
                        # 处理其他可能的空格实体
                        content_html = content_html.replace('&#160;', ' ')
                        content_html = content_html.replace('&#xA0;', ' ')
                        
                        # 完全解码所有HTML实体
                        content_html = html.unescape(content_html)
                        
                        print(f"处理前文本：{contract.content[:100]}")
                        print(f"处理后文本：{content_html[:100]}")
                        
                        # 用BeautifulSoup解析处理后的HTML
                        soup = BeautifulSoup(content_html, 'html.parser', from_encoding='utf-8')
                        
                        # 处理段落
                        for p in soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                            # 获取文本内容 - 使用get_text保留空格
                            text = p.get_text(separator=' ', strip=False)
                            
                            # 再次处理可能的HTML实体
                            original_text = text
                            text = text.replace('&nbsp;', ' ')
                            text = text.replace('&#160;', ' ')
                            text = text.replace('&#xA0;', ' ')
                            
                            # 处理其他非标准空格字符
                            text = text.replace('\u00A0', ' ')  # 不换行空格
                            text = text.replace('\u2002', ' ')  # 半角空格
                            text = text.replace('\u2003', ' ')  # 全角空格
                            text = text.replace('\u2004', ' ')  # 三分之一空格
                            text = text.replace('\u2005', ' ')  # 四分之一空格
                            text = text.replace('\u2006', ' ')  # 六分之一空格
                            text = text.replace('\u2007', ' ')  # 数字空格
                            text = text.replace('\u2008', ' ')  # 标点空格
                            text = text.replace('\u2009', ' ')  # 窄空格
                            text = text.replace('\u200A', ' ')  # 头发空格
                            text = text.replace('\u202F', ' ')  # 窄不换行空格
                            text = text.replace('\u205F', ' ')  # 中度数学空格
                            
                            text = html.unescape(text)  # 解码所有HTML实体
                            
                            if not text.strip():
                                continue
                                
                            # 创建段落
                            paragraph = document.add_paragraph()
                            
                            # 检查文本是否包含复选框字符
                            if '□' in text or '\u2610' in text or '\u25A1' in text:
                                # 使用特殊处理函数处理包含复选框的文本
                                handle_checkbox(paragraph, text, p)
                                
                                # 不再需要进一步处理段落内容
                                continue
                            
                            # 处理标题样式
                            if p.name and p.name.startswith('h') and len(p.name) == 2:
                                try:
                                    # 提取标题级别 (1-6)
                                    level = int(p.name[1])
                                    if 1 <= level <= 6:
                                        # 设置标题样式
                                        paragraph.style = f'Heading {level}'
                                except ValueError:
                                    pass  # 不是有效的标题标签
                            
                            # 应用段落样式
                            apply_paragraph_style(paragraph, p)
                            
                            # 处理段落中的格式
                            for content in p.contents:
                                if isinstance(content, str):
                                    paragraph.add_run(content)
                                elif content.name in ['strong', 'b']:
                                    run = paragraph.add_run(content.get_text())
                                    run.bold = True
                                elif content.name in ['em', 'i']:
                                    run = paragraph.add_run(content.get_text())
                                    run.italic = True
                                elif content.name == 'u':
                                    run = paragraph.add_run(content.get_text())
                                    run.underline = True
                                elif content.name == 'span':
                                    run = paragraph.add_run(content.get_text())
                                    
                                    # 处理span的样式
                                    if content.get('style'):
                                        style = content.get('style').lower()
                                        
                                        # 处理字体加粗
                                        if 'font-weight:bold' in style or 'font-weight: bold' in style:
                                            run.bold = True
                                        
                                        # 处理斜体
                                        if 'font-style:italic' in style or 'font-style: italic' in style:
                                            run.italic = True
                                        
                                        # 处理下划线
                                        if 'text-decoration:underline' in style or 'text-decoration: underline' in style:
                                            run.underline = True
                                        
                                        # 处理字体大小 (单位是pt)
                                        font_size_match = re.search(r'font-size:\s*(\d+)pt', style)
                                        if font_size_match:
                                            size = int(font_size_match.group(1))
                                            run.font.size = Pt(size)
                                            
                                        # 处理字体颜色 (使用颜色名称或十六进制值)
                                        color_match = re.search(r'color:\s*(#[0-9a-fA-F]{6}|[a-zA-Z]+)', style)
                                        if color_match:
                                            color_value = color_match.group(1).lower()
                                            # 处理常见颜色名称
                                            if color_value == 'red':
                                                run.font.color.rgb = RGBColor(255, 0, 0)
                                            elif color_value == 'green':
                                                run.font.color.rgb = RGBColor(0, 128, 0)
                                            elif color_value == 'blue':
                                                run.font.color.rgb = RGBColor(0, 0, 255)
                                            elif color_value == 'black':
                                                run.font.color.rgb = RGBColor(0, 0, 0)
                                            elif color_value.startswith('#'):
                                                # 处理十六进制颜色值 (#RRGGBB)
                                                try:
                                                    hex_color = color_value.lstrip('#')
                                                    r = int(hex_color[0:2], 16)
                                                    g = int(hex_color[2:4], 16)
                                                    b = int(hex_color[4:6], 16)
                                                    run.font.color.rgb = RGBColor(r, g, b)
                                                except (ValueError, IndexError):
                                                    # 如果解析失败，使用默认颜色
                                                    pass
                                elif content.name == 'br':
                                    paragraph.add_run('\n')
                                else:
                                    paragraph.add_run(content.get_text())
                    else:
                        document.add_paragraph('合同内容为空')
                    
                    # 创建响应
                    file_stream = BytesIO()
                    document.save(file_stream)
                    file_stream.seek(0)
                    
                    # 创建操作记录
                    ContractAction.objects.create(
                        contract=contract,
                        user=request.user,
                        action_type='download',
                        description='下载了合同Word文档'
                    )
                    
                    # 设置响应头
                    response = HttpResponse(
                        file_stream.getvalue(),
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    response['Content-Disposition'] = f'attachment; filename="{contract.number or "contract"}_{contract.title or "untitled"}.docx"'
                    
                    print("Word文档生成成功，准备发送下载响应")
                    return response
                    
                except Exception as doc_error:
                    print(f"Word文档生成过程中出错: {str(doc_error)}")
                    import traceback
                    traceback.print_exc()
                    return Response(
                        {"error": f"Word文档生成失败: {str(doc_error)}"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
            else:
                print(f"不支持的文件格式: {file_format}")
                return Response(
                    {"error": f"不支持的文件格式: {file_format}"},
                    status=status.HTTP_400_BAD_REQUEST
                )
                
        except Exception as e:
            print(f"下载合同过程中出现未处理异常: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response(
                {"error": f"下载合同失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['get'])
    def template_suggestions(self, request):
        """根据筛选条件获取推荐模板"""
        contract_type = request.query_params.get('contract_type', '')
        industry = request.query_params.get('industry', '')
        scene = request.query_params.get('scene', '')
        
        # 构建查询
        queryset = ContractTemplate.objects.filter(is_active=True)
        
        # 应用筛选条件
        if contract_type:
            queryset = queryset.filter(contract_type=contract_type)
        if industry:
            queryset = queryset.filter(industry=industry)
        if scene:
            queryset = queryset.filter(scene=scene)
            
        # 限制返回数量
        templates = queryset[:10]
        
        serializer = ContractTemplateListSerializer(templates, many=True)
        return Response(serializer.data)

class RegulationQueryView(APIView):
    """法规查询视图"""
    permission_classes = [permissions.IsAuthenticated]
    
    def post(self, request):
        """处理法规查询请求"""
        try:
            # 获取请求数据
            issue_title = request.data.get('issue_title', '')
            issue_content = request.data.get('issue_content', '')
            issue_problem = request.data.get('issue_problem', '')
            contract_type = request.data.get('contract_type', '通用合同')
            
            # 预处理文本，替换HTML实体
            issue_title = process_html_entities(issue_title)
            issue_content = process_html_entities(issue_content)
            issue_problem = process_html_entities(issue_problem)
            contract_type = process_html_entities(contract_type)
            
            # 验证必要参数
            if not issue_title or not issue_problem:
                return Response(
                    {"error": "缺少必要参数"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 调用AI查询相关法规
            result = query_regulation({
                'issue_title': issue_title,
                'issue_content': issue_content,
                'issue_problem': issue_problem,
                'contract_type': contract_type
            })
            
            # 返回查询结果
            return Response(result, status=status.HTTP_200_OK)
            
        except Exception as e:
            return Response(
                {"error": f"法规查询失败: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            ) 